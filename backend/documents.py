import os
import subprocess
from docx import Document

def make_client_invoice(name, product, unit, price):
    document = Document()
    document.add_heading('Invoice', 0)

    # Greeting
    p1 = document.add_paragraph('Dear ')
    p1.add_run(name).bold = True
    p1.add_run(',')

    # Body
    p2 = document.add_paragraph("Please find the attached invoice for you: ")
    p2.add_run(str(unit)).bold = True
    p2.add_run(' units of ')
    p2.add_run(product).bold = True
    p2.add_run('.')

    [document.add_paragraph('') for _ in range(2)]

    # Table
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    headers = ["Product Name", "Units", "Unit Price", "Total Price"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True

    row_cells = table.add_row().cells
    row_cells[0].text = product
    row_cells[1].text = f"{unit:,.2f}"
    row_cells[2].text = f"{price:,.2f}"
    row_cells[3].text = f"{unit * price:,.2f}"

    [document.add_paragraph('') for _ in range(10)]

    document.add_paragraph("We appreciate your business and look forward to collaborating again!")
    document.add_paragraph("Sincerely,")
    document.add_paragraph("Jay")

    docx_path = f"{name}.docx"
    document.save(docx_path)
    return docx_path

def doc_to_pdf_mac(docx_path):
    docx_path = os.path.abspath(docx_path)
    output_dir = os.path.dirname(docx_path)
    subprocess.run([
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_dir,
        docx_path
    ], check=True)
    print(f"PDF saved to: {docx_path.replace('.docx', '.pdf')}")

# Example usage
if __name__ == "__main__":
    docx_file = make_client_invoice("Elon", "Falcon", 10, 12500000)
    doc_to_pdf_mac(docx_file)